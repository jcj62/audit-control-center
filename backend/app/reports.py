from __future__ import annotations

import os
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from sqlalchemy import text
from sqlalchemy.orm import Session

from .models import Audit

try:
    from docx import Document
    from docx.enum.section import WD_ORIENT
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
except ImportError:  # pragma: no cover
    Document = None
    WD_ORIENT = None
    WD_ROW_HEIGHT_RULE = None
    WD_ALIGN_PARAGRAPH = None
    OxmlElement = None
    qn = None
    Cm = None
    Pt = None


def set_calibri_font(paragraph, size: int = 11) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)


def set_cell_margins(cell, top_cm: float = 0.2, bottom_cm: float = 0.2, left_cm: float = 0.3, right_cm: float = 0.3) -> None:
    tc = cell._element
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")

    for name, value in [
        ("top", top_cm),
        ("bottom", bottom_cm),
        ("left", left_cm),
        ("right", right_cm),
    ]:
        node = OxmlElement(f"w:{name}")
        node.set(qn("w:w"), str(int(value * 567)))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)

    tc_pr.append(tc_mar)


def set_cell_border(cell) -> None:
    tc = cell._element
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")

    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:color"), "000000")
        tc_borders.append(element)

    tc_pr.append(tc_borders)


def insert_images_dynamic(cell, image_paths: list[str], max_images: int) -> None:
    if not image_paths:
        return

    cell.text = ""
    images_to_use = image_paths[:max_images]

    for image_path in images_to_use:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()

        try:
            run.add_picture(image_path, width=Cm(7.2), height=Cm(4))
        except Exception as exc:
            print(f"Image insert error: {exc}")


def insert_images_uniform(cell, image_paths: list[str]) -> None:
    if not image_paths:
        return

    cell.text = ""
    for image_path in image_paths:
        paragraph = cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()

        try:
            run.add_picture(image_path, width=Cm(7.0), height=Cm(4.0))
        except Exception as exc:
            print(f"Image insert error: {exc}")


def _build_grouped_faults(db: Session, audit_id: int, media_dir: Path) -> list[tuple[str, str, dict]]:
    faults = db.execute(
        text(
            """
            SELECT cluster_id, building, fault_type, message, image_path
            FROM faults
            WHERE audit_id = :audit_id
            ORDER BY cluster_id
            """
        ),
        {"audit_id": audit_id},
    ).fetchall()

    grouped_faults = defaultdict(
        lambda: {
            "building": "",
            "fault_type": "",
            "remarks": [],
            "locations": [],
            "images": [],
        }
    )

    for fault in faults:
        key = (fault.building, fault.fault_type)
        grouped_faults[key]["building"] = fault.building
        grouped_faults[key]["fault_type"] = fault.fault_type
        grouped_faults[key]["remarks"].append(fault.message)
        grouped_faults[key]["locations"].append(fault.building)

        if fault.image_path:
            full_path = os.path.join(str(media_dir), fault.image_path)
            if os.path.exists(full_path):
                grouped_faults[key]["images"].append(full_path)

    return [(building, fault_type, data) for (building, fault_type), data in grouped_faults.items()]


def _render_report_table(document, grouped_faults: list[tuple[str, str, dict]], uniform_images: bool = False) -> None:
    table = document.add_table(rows=1, cols=3)
    table.autofit = False

    widths = [1.25, 7.44, 7.75]
    table_xml = table._element
    table_properties = table_xml.tblPr

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    table_properties.append(layout)

    table_grid = table_xml.tblGrid
    for child in list(table_grid):
        table_grid.remove(child)

    for width in widths:
        grid_column = OxmlElement("w:gridCol")
        grid_column.set(qn("w:w"), str(int(width * 567)))
        table_grid.append(grid_column)

    for index, width in enumerate(widths):
        for row in table.rows:
            row.cells[index].width = Cm(width)

    header = table.rows[0]
    header.height = Cm(1.04)
    header.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY

    for index, header_text in enumerate(["Si No", "Images", "Remarks"]):
        cell = header.cells[index]
        cell.text = header_text
        set_cell_border(cell)
        set_cell_margins(cell)

        paragraph = cell.paragraphs[0]
        set_calibri_font(paragraph, 12)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    serial = 1
    for (_, fault_type, data) in grouped_faults:
        row_cells = table.add_row().cells

        for cell in row_cells:
            set_cell_border(cell)
            set_cell_margins(cell)

        row_cells[0].text = str(serial)
        set_calibri_font(row_cells[0].paragraphs[0])

        images = list(dict.fromkeys(data["images"]))
        remarks_list = list(dict.fromkeys(data["remarks"]))
        locations = list(dict.fromkeys(data["locations"]))

        if uniform_images:
            insert_images_uniform(row_cells[1], images)
        else:
            num_lines = len(remarks_list)
            if len(remarks_list) == 1:
                num_lines += 1

            max_images = max(1, num_lines // 2)
            max_images = min(max_images, 5)
            insert_images_dynamic(row_cells[1], images, max_images)

        remarks_cell = row_cells[2]
        fault_type_text = fault_type.upper()

        if len(remarks_list) == 1:
            paragraph = remarks_cell.paragraphs[0]
            paragraph.add_run(remarks_list[0] or "")
            set_calibri_font(paragraph)

            fault_paragraph = remarks_cell.add_paragraph()
            fault_paragraph.add_run(fault_type_text)
            set_calibri_font(fault_paragraph)
        else:
            paragraph = remarks_cell.paragraphs[0]
            paragraph.add_run(fault_type_text)
            set_calibri_font(paragraph)

            for remark in remarks_list:
                bullet = remarks_cell.add_paragraph()
                bullet.add_run(f"- {remark or ''}")
                set_calibri_font(bullet)

        location_paragraph = remarks_cell.add_paragraph()
        location_paragraph.add_run(", ".join(locations))
        set_calibri_font(location_paragraph)

        serial += 1


def _generate_docx_report(db: Session, audit: Audit, output_dir: Path, media_dir: Path, *, uniform_images: bool = False) -> Path:
    if any(item is None for item in [Document, WD_ORIENT, WD_ROW_HEIGHT_RULE, WD_ALIGN_PARAGRAPH, OxmlElement, qn, Cm, Pt]):
        raise RuntimeError("python-docx is not installed")

    output_dir.mkdir(parents=True, exist_ok=True)
    safe_name = audit.audit_name.strip().replace(" ", "_").lower()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    suffix = "_uniform" if uniform_images else ""
    file_name = f"{safe_name}_{audit.id}{suffix}_{timestamp}.docx"
    file_path = output_dir / file_name

    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width

    title_text = "Electrical Safety Audit Report"
    if uniform_images:
        title_text = "Electrical Safety Audit Report - Uniform Images"
    title = document.add_heading(title_text, 0)
    set_calibri_font(title)
    document.add_paragraph("")

    grouped_faults = _build_grouped_faults(db, audit.id, media_dir)

    if not grouped_faults:
        paragraph = document.add_paragraph("No faults recorded.")
        set_calibri_font(paragraph)
        document.save(file_path)
        return file_path

    _render_report_table(document, grouped_faults, uniform_images=uniform_images)
    document.save(file_path)
    print(f"Report generated: {file_path}")
    return file_path


def generate_docx_report(db: Session, audit: Audit, output_dir: Path, media_dir: Path) -> Path:
    return _generate_docx_report(db, audit, output_dir, media_dir, uniform_images=False)


def generate_docx_report_uniform_images(db: Session, audit: Audit, output_dir: Path, media_dir: Path) -> Path:
    return _generate_docx_report(db, audit, output_dir, media_dir, uniform_images=True)
